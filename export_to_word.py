import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pygments import lex
from pygments.lexers import PythonLexer
from pygments.style import Style
from pygments.token import Token, Keyword, Name, Comment, String, Number, Operator, Punctuation, Text
from pygments.styles.monokai import MonokaiStyle


def set_paragraph_shading(paragraph, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    paragraph.paragraph_format.element.get_or_add_pPr().append(shading)


def python_to_word(python_file, output_docx, show_line_numbers=False):
    if not os.path.exists(python_file):
        print(f"错误：Python文件 {python_file} 不存在！")
        return

    with open(python_file, 'r', encoding='utf-8') as f:
        code_content = f.read()

    doc = Document()

    title = doc.add_heading(python_file, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"文件路径: {os.path.abspath(python_file)}")
    doc.add_paragraph(f"代码行数: {len(code_content.splitlines())} 行")
    line_status = "开启" if show_line_numbers else "关闭"
    doc.add_paragraph(f"行号显示: {line_status}")
    doc.add_paragraph()

    style = MonokaiStyle
    style_map = {}
    for ttype, tstyle in style:
        if tstyle and tstyle['color']:
            hex_color = tstyle['color']
            if len(hex_color) == 3:
                hex_color = ''.join([c * 2 for c in hex_color])
            style_map[ttype] = RGBColor(
                int(hex_color[0:2], 16),
                int(hex_color[2:4], 16),
                int(hex_color[4:6], 16)
            )

    line_number_color = RGBColor(144, 144, 144)

    if show_line_numbers:
        code_lines = code_content.splitlines(keepends=True)
        total_lines = len(code_lines)
        line_num_width = len(str(total_lines))

        for line_idx, line in enumerate(code_lines):
            line_num = str(line_idx + 1).rjust(line_num_width) + "  "

            code_block = doc.add_paragraph()
            code_block.paragraph_format.space_before = Pt(0)
            code_block.paragraph_format.space_after = Pt(0)
            code_block.paragraph_format.line_spacing = 1.0
            set_paragraph_shading(code_block, '272822')

            line_num_run = code_block.add_run(line_num)
            line_num_run.font.name = 'Courier New'
            line_num_run.font.size = Pt(10)
            line_num_run.font.color.rgb = line_number_color

            lexer = PythonLexer()
            tokens = lex(line, lexer)

            for ttype, value in tokens:
                if not value:
                    continue

                run = code_block.add_run(value)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)

                current = ttype
                while current:
                    if current in style_map:
                        run.font.color.rgb = style_map[current]
                        break
                    current = current.parent
    else:
        code_block = doc.add_paragraph()
        code_block.paragraph_format.space_before = Pt(0)
        code_block.paragraph_format.space_after = Pt(0)
        code_block.paragraph_format.line_spacing = 1.0
        set_paragraph_shading(code_block, '272822')

        lexer = PythonLexer()
        tokens = lex(code_content, lexer)

        for ttype, value in tokens:
            if not value:
                continue

            run = code_block.add_run(value)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

            current = ttype
            while current:
                if current in style_map:
                    run.font.color.rgb = style_map[current]
                    break
                current = current.parent

    doc.save(output_docx)
    print(f"Word文档已成功保存至: {output_docx}")


if __name__ == "__main__":
    python_file = "create_ppt.py"
    output_docx = "create_ppt_v2.docx"
    show_line_numbers = False

    python_to_word(python_file, output_docx, show_line_numbers=show_line_numbers)
